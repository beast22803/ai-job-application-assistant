
"""
analyzer.py — Parsing documents (PDF/DOCX) and performing AI-powered
job analysis, resume extraction, ATS matching, and suggestion generation.
"""

from __future__ import annotations

import io
import json
import re
import os
from typing import BinaryIO
from datetime import datetime

from pypdf import PdfReader
from docx import Document
from langchain_core.messages import HumanMessage, SystemMessage

from app.core.config import get_llm

# ── Document Parsers ──────────────────────────────────────────────────────────

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract plain text from PDF file bytes."""
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()
    except Exception as e:
        raise ValueError(f"Failed to parse PDF resume: {e}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract plain text from Word (.docx) file bytes."""
    try:
        doc = Document(io.BytesIO(file_bytes))
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        return "\n".join(text).strip()
    except Exception as e:
        raise ValueError(f"Failed to parse Word document: {e}")


def extract_resume_text(filename: str, content: bytes) -> str:
    """Helper to detect file type and extract text content."""
    lower_fn = filename.lower()
    if lower_fn.endswith(".pdf"):
        return extract_text_from_pdf(content)
    elif lower_fn.endswith(".docx"):
        return extract_text_from_docx(content)
    else:
        # Fallback to plain text decoding
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            return content.decode("latin-1")


# ── AI Analysis Helpers ───────────────────────────────────────────────────────

def load_prompt(filename: str) -> str:
    """Load prompt content from a file in the prompts directory."""
    dir_path = os.path.dirname(os.path.abspath(__file__))
    path = os.path.join(dir_path, "prompts", filename)
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def parse_date(date_str: str) -> datetime | None:
    """Robustly parse MM/YYYY, textual Month YYYY, or YYYY date strings."""
    if not date_str:
        return None
    date_str = date_str.lower().strip()
    if date_str in ('present', 'current', 'now', ''):
        return datetime.now()
    
    # 1. Try MM/YYYY
    mmyyyy = re.search(r'(\d{1,2})/(\d{4})', date_str)
    if mmyyyy:
        return datetime(int(mmyyyy.group(2)), int(mmyyyy.group(1)), 1)
        
    # 2. Try MM-YYYY
    mmyyyy_dash = re.search(r'(\d{1,2})-(\d{4})', date_str)
    if mmyyyy_dash:
        return datetime(int(mmyyyy_dash.group(2)), int(mmyyyy_dash.group(1)), 1)
        
    # 3. Try textual month and 4-digit year (e.g. "january 2023", "jan 2023")
    months_map = {
        'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
        'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12,
        'january': 1, 'february': 2, 'march': 3, 'april': 4, 'june': 6,
        'july': 7, 'august': 8, 'september': 9, 'october': 10, 'november': 11, 'december': 12
    }
    month_match = re.search(r'([a-z]+)\s+(\d{4})', date_str)
    if month_match:
        m_name = month_match.group(1)
        y_val = int(month_match.group(2))
        if m_name in months_map:
            return datetime(y_val, months_map[m_name], 1)
            
    # 4. Try 4-digit year anywhere
    yyyy = re.search(r'\b(\d{4})\b', date_str)
    if yyyy:
        return datetime(int(yyyy.group(1)), 1, 1)
        
    return None


def calculate_experience_months(experience_list: list[dict]) -> int:
    """Sum total experience duration in months from experience list."""
    total_months = 0
    for exp in experience_list:
        if not isinstance(exp, dict):
            continue
        start_str = str(exp.get("start") or exp.get("startDate") or "").lower().strip()
        end_str = str(exp.get("end") or exp.get("endDate") or "present").lower().strip()
        
        start_date = parse_date(start_str)
        end_date = parse_date(end_str)
        
        if start_date and end_date:
            months = (end_date.year - start_date.year) * 12 + (end_date.month - start_date.month)
            if months > 0:
                total_months += months
    return total_months


# ── AI Analysis Functions ─────────────────────────────────────────────────────

def analyze_job_description(job_text: str) -> dict:
    """Extract key details from a job description using the LLM."""
    llm = get_llm()
    system_prompt = load_prompt("prompt_job_analyzer.txt")
    system_prompt = system_prompt.replace("{{ $json.text }}", job_text)

    response = llm.invoke([
        SystemMessage(content=system_prompt),
        HumanMessage(content="Analyze this job description.")
    ])

    cleaned_content = _clean_json_output(response.content)
    try:
        data = json.loads(cleaned_content)
        # Ensure all required keys exist (both old and new schemas)
        for k in ["role", "company", "job_level", "employment_type", "industry", "location", "experience_required",
                  "education_required", "skills_required", "skills_preferred", "tools_and_technologies",
                  "responsibilities", "keywords", "soft_skills"]:
            if k not in data:
                data[k] = [] if "skills" in k or "keywords" in k or k in ["tools_and_technologies", "responsibilities", "soft_skills"] else ""
        
        # Map new keys to old keys for compatibility
        data["role_information"] = data.get("role", "")
        data["experience_requirements"] = data.get("experience_required", "")
        data["required_skills"] = data.get("skills_required", [])
        data["preferred_skills"] = data.get("skills_preferred", [])
        data["ats_keywords"] = data.get("keywords", [])
        if "company" not in data or not data["company"]:
            data["company"] = "Target Company"
        return data
    except Exception as e:
        print(f"Error parsing job analysis JSON: {e}. Raw: {response.content}")
        return {
            "role": "Software Engineer",
            "job_level": "Unknown",
            "employment_type": "Full-Time",
            "industry": "Unknown",
            "location": "",
            "experience_required": "Not specified",
            "education_required": "",
            "skills_required": ["React", "TypeScript"],
            "skills_preferred": [],
            "tools_and_technologies": [],
            "responsibilities": [],
            "keywords": ["React", "TypeScript"],
            "soft_skills": [],
            "role_information": "Software Engineer",
            "experience_requirements": "Not specified",
            "required_skills": ["React", "TypeScript"],
            "preferred_skills": [],
            "ats_keywords": ["React", "TypeScript"],
            "company": "Target Company"
        }


def normalize_parsed_resume(data: dict) -> dict:
    """Normalize and map common LLM parser schema variations to the expected schema."""
    if not isinstance(data, dict):
        return {}

    # Map root keys
    if "workExperience" in data and ("experience" not in data or not data["experience"]):
        data["experience"] = data.pop("workExperience")
    elif "work_experience" in data and ("experience" not in data or not data["experience"]):
        data["experience"] = data.pop("work_experience")

    # Map experience items
    if isinstance(data.get("experience"), list):
        normalized_exp = []
        for exp in data["experience"]:
            if not isinstance(exp, dict):
                continue
            item = {}
            # Map title
            item["title"] = exp.get("title") or exp.get("job_title") or exp.get("jobTitle") or ""
            # Map company
            item["company"] = exp.get("company") or exp.get("employer") or exp.get("organization") or ""
            
            # Map start and end dates
            start = str(exp.get("start") or "").strip()
            end = str(exp.get("end") or "").strip()
            dates = exp.get("dates") or exp.get("date_range") or exp.get("dateRange") or exp.get("year") or ""
            
            range_str = ""
            if any(delim in start for delim in ("-", "–", "—", " to ")):
                range_str = start
            elif dates:
                range_str = str(dates).strip()
                
            if range_str:
                parts = re.split(r'\s+to\s+|\s*[-–—]\s*', range_str, flags=re.IGNORECASE)
                if len(parts) >= 2:
                    start = parts[0].strip()
                    end = parts[1].strip()
                else:
                    start = range_str
                    end = ""
                    
            if start.lower() == "none":
                start = ""
            if end.lower() == "none":
                end = ""
            
            item["start"] = start
            item["end"] = end
            
            # Map description
            desc = exp.get("description")
            if not desc:
                # Check for "experience" or "bullets" or "responsibilities"
                val = exp.get("experience") or exp.get("bullets") or exp.get("responsibilities") or exp.get("achievements") or ""
                if isinstance(val, list):
                    desc = "\n".join([str(x) for x in val])
                else:
                    desc = str(val)
            item["description"] = str(desc or "").strip()
            normalized_exp.append(item)
        data["experience"] = normalized_exp

    # Map projects items
    if isinstance(data.get("projects"), list):
        normalized_proj = []
        for proj in data["projects"]:
            if not isinstance(proj, dict):
                continue
            item = {}
            item["name"] = proj.get("name") or proj.get("title") or proj.get("projectName") or ""
            item["technologies"] = proj.get("technologies") or proj.get("tech_stack") or proj.get("tech") or []
            if isinstance(item["technologies"], str):
                item["technologies"] = [t.strip() for t in item["technologies"].split(",")]
            
            desc = proj.get("description") or proj.get("highlights") or proj.get("bullets") or ""
            if isinstance(desc, list):
                desc = "\n".join([str(x) for x in desc])
            item["description"] = str(desc).strip()
            normalized_proj.append(item)
        data["projects"] = normalized_proj

    # Map education items
    if isinstance(data.get("education"), list):
        normalized_edu = []
        for edu in data["education"]:
            if not isinstance(edu, dict):
                continue
            item = {}
            item["degree"] = edu.get("degree") or edu.get("qualification") or ""
            item["institution"] = edu.get("institution") or edu.get("school") or edu.get("university") or ""
            
            year = edu.get("year") or edu.get("dates") or edu.get("graduation_year") or edu.get("year_range") or ""
            item["year"] = str(year).strip()
            normalized_edu.append(item)
        data["education"] = normalized_edu

    return data


def parse_optimized_resume_data(resume_data: dict, original_analysis: dict) -> dict:
    """
    Parses the generated resume_data JSON back into a structured resume_analysis dict,
    avoiding the second LLM call entirely.
    """
    header = resume_data.get("header", "")
    header_parts = [p.strip() for p in header.split("|") if p.strip()]
    
    # Defaults from original_analysis
    name = original_analysis.get("name") or original_analysis.get("candidate_name") or "Candidate"
    email = original_analysis.get("email", "")
    phone = original_analysis.get("phone", "")
    location = original_analysis.get("location", "")
    linkedin = original_analysis.get("linkedin", "")
    portfolio = original_analysis.get("portfolio", "")
    
    # Refine name and contact info using parts from the header if they match typical patterns
    if header_parts:
        # First part is name, unless it is a generic/fallback string
        first_part = header_parts[0]
        if first_part and first_part.lower() != "candidate":
            name = first_part
            
        for part in header_parts[1:]:
            if "@" in part:
                email = part
            elif "linkedin.com" in part or "linkedin" in part:
                linkedin = part
            elif "github.com" in part or "github" in part or "portfolio" in part:
                portfolio = part
            elif any(c.isdigit() for c in part) and len([c for c in part if c.isdigit()]) >= 7:
                phone = part
            else:
                if part and not location:
                    location = part

    analysis = {
        "name": name,
        "email": email,
        "phone": phone,
        "location": location,
        "linkedin": linkedin,
        "portfolio": portfolio,
        "certifications": original_analysis.get("certifications", []),
        "languages": original_analysis.get("languages", []),
        "summary": resume_data.get("summary", ""),
        "skills": resume_data.get("skills", []),
        "experience": [],
        "projects": [],
        "education": []
    }
            
    # Parse experience
    # Format: "[Job Title] · [Company] · [Start]–[End]\n[bullet1]\n[bullet2]..."
    experience_raw = resume_data.get("experience", [])
    experience_structured = []
    experience_strings = []
    for exp_str in experience_raw:
        if not isinstance(exp_str, str) or not exp_str.strip():
            continue
        lines = [line.strip() for line in exp_str.split("\n") if line.strip()]
        if not lines:
            continue
        meta_line = lines[0]
        # Split meta_line by · or | or -
        meta_parts = [p.strip() for p in meta_line.split("·")]
        if len(meta_parts) < 3:
            meta_parts = [p.strip() for p in meta_line.split("|")]
        if len(meta_parts) < 3:
            meta_parts = [p.strip() for p in meta_line.split(" - ")]
            
        title = ""
        company = ""
        start = ""
        end = ""
        
        if len(meta_parts) >= 3:
            title = meta_parts[0]
            company = meta_parts[1]
            dates_str = " · ".join(meta_parts[2:])
            # Split dates by - or – or — or to
            date_parts = re.split(r'\s+to\s+|\s*[-–—]\s*', dates_str, flags=re.IGNORECASE)
            if len(date_parts) >= 2:
                start = date_parts[0].strip()
                end = date_parts[1].strip()
            else:
                start = dates_str
                end = ""
        elif len(meta_parts) == 2:
            title = meta_parts[0]
            company = meta_parts[1]
        else:
            title = meta_line
            
        description = "\n".join(lines[1:]).strip()
        
        experience_structured.append({
            "title": title,
            "company": company,
            "start": start,
            "end": end,
            "description": description
        })
        experience_strings.append(exp_str)
        
    analysis["experience_structured"] = experience_structured
    analysis["experience"] = experience_strings
    
    # Parse projects
    # Format: "[Project Name] — [tech1, tech2, tech3]\n[1–2 sentence description]"
    projects_raw = resume_data.get("projects", [])
    projects_structured = []
    projects_strings = []
    for proj_str in projects_raw:
        if not isinstance(proj_str, str) or not proj_str.strip():
            continue
        lines = [line.strip() for line in proj_str.split("\n") if line.strip()]
        if not lines:
            continue
        meta_line = lines[0]
        meta_parts = [p.strip() for p in meta_line.split("—")]
        if len(meta_parts) < 2:
            meta_parts = [p.strip() for p in meta_line.split(" - ")]
        if len(meta_parts) < 2:
            meta_parts = [p.strip() for p in meta_line.split("|")]
            
        name = ""
        tech = []
        if len(meta_parts) >= 2:
            name = meta_parts[0]
            tech = [t.strip() for t in meta_parts[1].split(",") if t.strip()]
        else:
            name = meta_line
            
        description = "\n".join(lines[1:]).strip()
        
        projects_structured.append({
            "name": name,
            "technologies": tech,
            "description": description
        })
        projects_strings.append(proj_str)
        
    analysis["projects_structured"] = projects_structured
    analysis["projects"] = projects_strings
    
    # Parse education
    # Format: "[Degree] · [Institution] · [Year]"
    education_raw = resume_data.get("education", [])
    education_structured = []
    education_strings = []
    for edu_str in education_raw:
        if not isinstance(edu_str, str) or not edu_str.strip():
            continue
        edu_parts = [p.strip() for p in edu_str.split("·")]
        if len(edu_parts) < 2:
            edu_parts = [p.strip() for p in edu_str.split("|")]
        if len(edu_parts) < 2:
            edu_parts = [p.strip() for p in edu_str.split(" - ")]
            
        degree = ""
        institution = ""
        year = ""
        
        if len(edu_parts) >= 3:
            degree = edu_parts[0]
            institution = edu_parts[1]
            year = " · ".join(edu_parts[2:])
        elif len(edu_parts) == 2:
            degree = edu_parts[0]
            institution = edu_parts[1]
        else:
            degree = edu_str
            
        education_structured.append({
            "degree": degree,
            "institution": institution,
            "year": year
        })
        education_strings.append(edu_str)
        
    analysis["education_structured"] = education_structured
    analysis["education"] = education_strings
    
    # Legacy compat
    analysis["candidate_name"] = analysis.get("name", "")
    
    return analysis


def analyze_resume(resume_text: str) -> dict:
    """Extract candidate details from a resume, strictly adhering to the facts in the text."""
    llm = get_llm()
    system_prompt = load_prompt("prompt_resume_analyzer.txt")
    system_prompt = system_prompt.replace("{{ $json.text }}", resume_text)

    response = llm.invoke([
        SystemMessage(content=system_prompt),
        HumanMessage(content="Extract info from this resume. You must return ONLY a raw JSON object matching the required schema. Do NOT return markdown lists, code blocks, or conversational text under any circumstances.")
    ])

    cleaned_content = _clean_json_output(response.content)
    try:
        data = json.loads(cleaned_content)
        data = normalize_parsed_resume(data)
        # Check all required new keys
        new_keys = ["name", "email", "phone", "location", "linkedin", "portfolio", "summary", "skills",
                    "experience", "projects", "education", "certifications", "languages"]
        for k in new_keys:
            if k not in data:
                data[k] = [] if k in ["skills", "experience", "projects", "education", "certifications", "languages"] else ""

        # Map to old keys for compatibility
        data["candidate_name"] = data.get("name", "")
        # Keep structured lists as is for downstream engines (e.g. generator)
        data["experience_structured"] = data.get("experience", [])
        data["projects_structured"] = data.get("projects", [])
        data["education_structured"] = data.get("education", [])

        # Map experience, projects, education to string lists for safety
        data["experience"] = [
            f"{exp.get('title', '')} at {exp.get('company', '')} ({exp.get('start', '')} - {exp.get('end', '')}): {exp.get('description', '')}"
            for exp in data.get("experience", [])
            if isinstance(exp, dict)
        ]
        data["projects"] = [
            f"{proj.get('name', '')} ({', '.join(proj.get('technologies', []))}): {proj.get('description', '')}"
            for proj in data.get("projects", [])
            if isinstance(proj, dict)
        ]
        data["education"] = [
            f"{edu.get('degree', '')} at {edu.get('institution', '')} ({edu.get('year', '')})"
            for edu in data.get("education", [])
            if isinstance(edu, dict)
        ]
        return data
    except Exception as e:
        print(f"Error parsing resume analysis JSON: {e}. Raw: {response.content}")
        return {
            "name": "Candidate",
            "email": "",
            "phone": "",
            "location": "",
            "linkedin": "",
            "portfolio": "",
            "summary": "",
            "skills": [],
            "experience": [],
            "projects": [],
            "education": [],
            "certifications": [],
            "languages": [],
            "candidate_name": "Candidate",
            "experience_structured": [],
            "projects_structured": [],
            "education_structured": []
        }


def run_ats_scoring(job_analysis: dict, resume_analysis: dict, raw_job: str, raw_resume: str) -> dict:
    """
    Compute deterministic matching and then prompt the ATS Matching Engine (prompt_ats_matching_engine.txt)
    to calculate the final score and other breakdown categories.
    """
    req_skills = [str(s).strip() for s in job_analysis.get("skills_required", [])]
    res_skills = [str(s).strip() for s in resume_analysis.get("skills", [])]

    # BR-8: Deduplicate resume skills
    seen_skills = set()
    deduped_res_skills = []
    for skill in res_skills:
        norm = re.sub(r'[^a-z0-9]', '', skill.lower())
        if norm not in seen_skills:
            seen_skills.add(norm)
            deduped_res_skills.append(skill)
    resume_analysis["skills"] = deduped_res_skills

    # BR-4: Deterministic skill matching
    matched_skills = []
    missing_skills = []

    for req in req_skills:
        req_lower = req.lower().strip()
        req_norm = re.sub(r'[^a-z0-9]', '', req_lower)
        found = False
        for cs in deduped_res_skills:
            cs_lower = cs.lower().strip()
            cs_norm = re.sub(r'[^a-z0-9]', '', cs_lower)
            if cs_lower == req_lower or req_lower in cs_lower or cs_lower in req_lower or cs_norm == req_norm:
                found = True
                break
        if found:
            matched_skills.append(req)
        else:
            missing_skills.append(req)

    deterministic_skill_score = round((len(matched_skills) / len(req_skills)) * 100) if req_skills else 0

    # BR-5: Experience Calculation
    exp_list = resume_analysis.get("experience_structured", [])
    if not exp_list:
        # Check if experience is list of dicts
        exp_list = [e for e in resume_analysis.get("experience", []) if isinstance(e, dict)]

    total_months = calculate_experience_months(exp_list)
    total_years = round(total_months / 12.0, 1)

    req_years_match = re.search(r'(\d+)', str(job_analysis.get("experience_required", "")))
    required_years = int(req_years_match.group(1)) if req_years_match else 0
    experience_met = total_years >= required_years

    # BR-5b: Deterministic experience match score
    if required_years == 0:
        experience_match_score = 85
    else:
        ratio = total_years / required_years
        if ratio >= 1.5:
            experience_match_score = 95
        elif ratio >= 1.0:
            experience_match_score = 85
        elif ratio >= 0.8:
            experience_match_score = 72
        elif ratio >= 0.6:
            experience_match_score = 60
        elif ratio >= 0.4:
            experience_match_score = 45
        else:
            experience_match_score = 30

    # BR-9: Keyword Density Check
    keywords = job_analysis.get("keywords") or job_analysis.get("skills_required") or []
    resume_json_str = json.dumps(resume_analysis).lower()
    keyword_density = {}
    low_density_keywords = []
    for kw in keywords:
        kw_lower = str(kw).lower()
        escaped = re.escape(kw_lower)
        matches = re.findall(escaped, resume_json_str)
        count = len(matches)
        keyword_density[kw] = count
        if count < 2:
            low_density_keywords.append(kw)

    deterministic = {
        "matchedSkills": matched_skills,
        "missingSkills": missing_skills,
        "deterministicSkillScore": deterministic_skill_score,
        "experienceMatchScore": experience_match_score,
        "totalExperienceMonths": total_months,
        "totalExperienceYears": total_years,
        "requiredYears": required_years,
        "experienceMet": experience_met,
        "keywordDensity": keyword_density,
        "lowDensityKeywords": low_density_keywords
    }

    # Now load and render prompt_ats_matching_engine.txt
    llm = get_llm()
    system_prompt = load_prompt("prompt_ats_matching_engine.txt")
    system_prompt = system_prompt.replace("{{ JSON.stringify($json.deterministic, null, 2) }}", json.dumps(deterministic, indent=2))
    system_prompt = system_prompt.replace("{{ JSON.stringify($json.job, null, 2) }}", json.dumps(job_analysis, indent=2))
    system_prompt = system_prompt.replace("{{ JSON.stringify($json.resume, null, 2) }}", json.dumps(resume_analysis, indent=2))

    response = llm.invoke([
        SystemMessage(content=system_prompt),
        HumanMessage(content="Evaluate the resume and return scoring details.")
    ])

    cleaned_content = _clean_json_output(response.content)
    try:
        data = json.loads(cleaned_content)
        if not isinstance(data, dict):
            print(f"Warning: ATS engine returned non-dict: {data}")
            data = {}
    except Exception as e:
        print(f"Error parsing ATS matching engine JSON: {e}. Raw: {response.content}")
        data = {}

    breakdown = data.get("breakdown", {})
    skill_pct = deterministic_skill_score
    exp_score = experience_match_score
    sem_score = breakdown.get("semantic_match", 70)
    fmt_score = breakdown.get("formatting", 90)
    
    penalty = len(missing_skills) * 5

    # Compute final score mathematically in python to prevent LLM hallucinations/math errors
    final_score = round((skill_pct * 0.40) + (exp_score * 0.30) + (sem_score * 0.20) + (fmt_score * 0.10) - penalty)
    final_score = max(0, min(100, final_score))

    res = {
        "final_score": final_score,
        "missing_skills": missing_skills,
        "matched_skills": matched_skills,
        "scores": {
            "skill_match": skill_pct,
            "experience_match": exp_score,
            "semantic_match": sem_score,
            "formatting_score": fmt_score,
            "penalty": penalty
        },
        "summary": data.get("summary", "Resume evaluated successfully."),
        "deterministic": deterministic
    }
    return res


def generate_suggestions(job_analysis: dict, resume_analysis: dict, ats_results: dict) -> list[str]:
    """Generate 5-8 actionable, job-specific resume improvements based on ATS gaps using prompt_suggestions_engine.txt."""
    llm = get_llm()
    system_prompt = load_prompt("prompt_suggestions_engine.txt")
    
    deterministic = ats_results.get("deterministic", {})
    ats_score = ats_results.get("final_score", 0)
    
    # Verdict computation matching server.py
    if ats_score >= 85:
        verdict = "Strong Apply"
    elif ats_score >= 70:
        verdict = "Apply"
    elif ats_score >= 55:
        verdict = "Apply with Resume Changes"
    else:
        verdict = "Skip"

    system_prompt = system_prompt.replace("{{ JSON.stringify($json.output, null, 2) }}", json.dumps(ats_results, indent=2))
    system_prompt = system_prompt.replace("{{ JSON.stringify($json.deterministic, null, 2) }}", json.dumps(deterministic, indent=2))
    system_prompt = system_prompt.replace("{{ $json.verdict }}", verdict)

    try:
        response = llm.invoke([
            SystemMessage(content=system_prompt),
            HumanMessage(content="Generate suggestions now.")
        ])
        data = json.loads(_clean_json_output(response.content))
        if not isinstance(data, dict):
            print(f"Warning: Suggestions engine returned non-dict: {data}")
            data = {}
        suggestions_list = data.get("resume_suggestions", [])
        
        # Extract the 'text' field of each suggestion to return as list[str]
        suggestions = [sug["text"] for sug in suggestions_list if isinstance(sug, dict) and "text" in sug]
        
        # Ensure we have between 5 and 8 suggestions
        suggestions = suggestions[:8]
        if len(suggestions) < 5:
            # Fallbacks
            defaults = [
                "Ensure resume sections are formatted with consistent title case",
                "Add quantitative metrics (e.g. percentages, dollars, hours saved) to past roles",
                "List skills in a dedicated, scanning-friendly column or grid",
                "Align job experience chronologically starting from the most recent",
                "Integrate missing skills explicitly in your project descriptions"
            ]
            for d in defaults:
                if len(suggestions) >= 5:
                    break
                if d not in suggestions:
                    suggestions.append(d)
        return suggestions
    except Exception as e:
        print(f"Error generating suggestions: {e}")
        return [
            "Add quantified achievements to experience bullets",
            "Highlight projects that match the target job stack",
            "List technical skills at the top of the resume",
            "Reorder experience chronologically starting from current role",
            "Ensure consistency in date formats (MM/YYYY) throughout the document"
        ]


# ── Utility Helpers ───────────────────────────────────────────────────────────

def _clean_json_output(text: str) -> str:
    """Helper to strip markdown JSON code block wrappers and conversational text surrounding the JSON object."""
    text = text.strip()
    
    # Extract only the JSON object if there is text surrounding it
    start_idx = text.find('{')
    end_idx = text.rfind('}')
    if start_idx != -1 and end_idx != -1 and end_idx > start_idx:
        text = text[start_idx:end_idx + 1]
        
    if text.startswith("```"):
        text = re.sub(r"^```(json)?\n", "", text)
        text = re.sub(r"\n```$", "", text)
    return text.strip()
